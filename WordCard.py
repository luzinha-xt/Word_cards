from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from unidecode import unidecode

lists_object = open('ss2-ss4_sentence_word_lists 2.txt', 'r')
lists = lists_object.read().split(',')

for list in lists:
  #### read word list
  file_object = open(list + '.txt', 'r')
  items = file_object.read().split(',')
  i = len(items)
  print(items, i)
  document = Document()
  #### set font
  style = document.styles['Normal']
  font = style.font
  font.name = 'Sofia Pro Soft'
  font.size = Pt(80)

  #### set Document/Margins
  for sec in document.sections:
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)

  ##### Table setting
  table = document.add_table(rows=i, cols=1, style='Table Grid')
  for col in table.columns:
      for cell in col.cells:
          for par in cell.paragraphs:
              par.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
  for row in table.rows:
      row.height = Pt(206)
  table.autofit = False

  def set_cell_border( _Cell, **kwargs):
    tc = _Cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
          tcBorders = OxmlElement('w:tcBorders')
          tcPr.append(tcBorders)

      # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
          edge_data = kwargs.get(edge)
          if edge_data:
              tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
              element = tcBorders.find(qn(tag))
              if element is None:
                  element = OxmlElement(tag)
                  tcBorders.append(element)

            # looks like order of attributes is important
              for key in ["sz", "val", "color", "space", "shadow"]:
                  if key in edge_data:
                      element.set(qn('w:{}'.format(key)), str(edge_data[key]))

#### Add table content
  num = 0 # type: int
  while num < i:
      unicode_text = items[num].decode("utf-8", "replace")
      clean_ascii = unidecode(unicode_text)
      table.cell(num, 0).text = clean_ascii
      table.cell(num, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
      table.cell(num, 0).paragraphs[0].paragraph_format.space_before =Pt(63)
      set_cell_border(
          table.cell(num, 0),
          top={"sz": 12,"val": "dashed"},
          bottom={"sz": 12,"val": "dashed"},
          start={"sz": 12,"val": "dashed"},
          end={"sz": 12,"val": "dashed"}, )
      num = num+1

  document.save(list + '.docx')
